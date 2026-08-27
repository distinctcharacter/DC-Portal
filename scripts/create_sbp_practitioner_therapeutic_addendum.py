from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "source-library" / "somatic-baseline"
OUTPUT_DIR = ROOT / "output" / "final-document-corrections" / "docx"
DOCX_PATH = OUTPUT_DIR / "Somatic_Baseline_Practitioner_Therapeutic_Addendum_v1.0.docx"
LOGO_PATH = ROOT / "public" / "assets" / "dc-logo.png"

DARK_PURPLE = "241239"
GOLD = "B58B2A"
BLACK = "111111"
SLATE = "374151"
WHITE = "FFFFFF"
WARM_FILL = "F8F5EC"
LIGHT_GOLD = "E6D8B7"
RULE = "C9B47A"
RISK = "7F1D1D"

TITLE = "Somatic Baseline Practitioner Therapeutic Addendum"
ASSET_ID = "DC-P01-SBP-TA01"


CONTENTS = [
    "Professional scope and qualification boundary",
    "Readiness, consent, and risk classification",
    "Practitioner implementation architecture",
    "Section-by-section clinical adjacency guidance",
    "Pacing, titration, and adverse-response decisions",
    "Relational, cultural, and systemic considerations",
    "Documentation, outcome review, and fidelity",
    "Printable practitioner tools",
    "Completion, transition, and references",
]

SCOPE_MATRIX = [
    ["Role", "Permitted use", "Boundary"],
    [
        "Licensed mental health or healthcare professional",
        "May integrate the protocol within the professional's existing scope, assessment process, treatment plan, documentation standards, and applicable law.",
        "The addendum does not expand licensure, replace clinical judgment, or establish a new treatment modality.",
    ],
    [
        "Coach, educator, consultant, or facilitator",
        "May support education, reflection, habit formation, environmental review, and non-clinical accountability.",
        "Must not diagnose, treat mental illness, represent the protocol as therapy, manage crises, or advise a client to disregard licensed care.",
    ],
    [
        "Distinct Character practitioner-layer user",
        "May use approved portal resources with clients who have appropriate access and within the user's verified professional role.",
        "Portal access is not a license, certification, clinical credential, or authorization to reproduce or teach the intellectual property.",
    ],
]

READINESS_GATE = [
    ["Review area", "Proceed when", "Pause, adapt, or refer when"],
    [
        "Immediate safety",
        "The client can maintain basic safety and participate voluntarily.",
        "There is imminent danger, self-harm risk, violence risk, severe impairment, or inability to maintain safety.",
    ],
    [
        "Current stability",
        "The client can orient to the present, communicate choices, and recover from mild activation with support.",
        "The client is acutely intoxicated, manic, psychotic, severely dissociated, medically unstable, or repeatedly unable to reorient.",
    ],
    [
        "Environment",
        "The work can occur without increasing material, relational, housing, workplace, or physical danger.",
        "Boundary changes or disclosure could provoke retaliation, coercion, abuse, financial harm, or loss of essential support.",
    ],
    [
        "Professional fit",
        "The client's needs fall within the practitioner's competence and role.",
        "The presentation requires trauma treatment, medical assessment, crisis response, substance-use treatment, or another service outside scope.",
    ],
    [
        "Consent and choice",
        "The client understands the educational purpose, can decline any exercise, and knows how to stop.",
        "Participation is coerced, confused with treatment, or tied to consequences that remove meaningful choice.",
    ],
]

RISK_LEVELS = [
    ["Level", "Observable conditions", "Practitioner response"],
    [
        "Green - workable engagement",
        "Mild activation; intact orientation; clear speech; client can choose, pause, and recover.",
        "Continue at a measured pace. Track state before and after. End with an ordinary-life next step.",
    ],
    [
        "Amber - reduced capacity",
        "Escalating distress, narrowing attention, confusion, marked shutdown, dizziness, panic sensations, or repeated loss of choice.",
        "Stop the exercise. Orient to the room, reduce intensity, return to neutral activity, and decide whether to end. Consult or refer when the pattern repeats.",
    ],
    [
        "Red - urgent safety concern",
        "Imminent self-harm or harm risk, medical emergency, severe disorganization, inability to maintain safety, or disclosure requiring mandated action.",
        "Follow emergency, safeguarding, reporting, and organizational procedures for the practitioner's jurisdiction and role. Do not continue protocol work.",
    ],
]

SESSION_ARCHITECTURE = [
    ["Stage", "Practitioner task", "Completion signal"],
    ["1. Frame", "Clarify purpose, limits, consent, confidentiality, and the client's right to stop.", "The client can explain what the session is and is not."],
    ["2. Baseline", "Observe current capacity using client language, ordinary behavior, and a brief 0-10 activation rating.", "The practitioner has a usable starting point without diagnosing from a proprietary score."],
    ["3. Select", "Choose one narrow target and the lowest sufficient intensity.", "The target is specific, present-focused, and inside scope."],
    ["4. Practice", "Guide one exercise or reflection while monitoring orientation, breath, speech, movement, choice, and distress.", "The client remains able to choose, describe, and pause."],
    ["5. Integrate", "Translate the observation into one decision, boundary, environmental change, or repeatable practice.", "The client leaves with a concrete next action."],
    ["6. Close", "Reassess state, document response, confirm support, and identify escalation criteria.", "The client is at or near baseline and knows what to do if distress rises later."],
]

SECTION_GUIDANCE = [
    ["Protocol element", "Practitioner use", "Interpretive limit"],
    [
        "Somatic Dysregulation Index (SDI)",
        "Use as a proprietary self-report conversation starter and repeated personal tracking tool.",
        "It is not a validated diagnostic instrument. Scores do not establish a disorder, trauma history, treatment need, or clinical severity.",
    ],
    [
        "Zone language",
        "Use Zones 1-3 as accessible descriptions of current capacity and behavioral options.",
        "Do not treat a zone label as a fixed nervous-system state, diagnosis, or objective physiological measurement.",
    ],
    [
        "Vagal and polyvagal language",
        "Use as an organizing lens when it helps the client understand regulation, connection, and recovery.",
        "Present polyvagal theory as influential and debated. Avoid claims that a behavior proves a specific vagal pathway or that a brief exercise mechanically resets the vagus nerve.",
    ],
    [
        "Environmental Friction Map",
        "Identify conditions that increase load and classify them as remove, restructure, regulate, or refer for specialized support.",
        "Do not individualize structural harm. Personal regulation is not a substitute for safety planning, workplace action, legal support, medical care, or material resources.",
    ],
    [
        "Tactical resets",
        "Offer as optional state-management practices. Start with the least activating option and monitor response.",
        "No technique is universally regulating. Stop for pain, dizziness, breathlessness, panic escalation, numbness, or disorientation.",
    ],
    [
        "Governance Log",
        "Track triggers, context, choices, recovery time, and functional outcomes to identify patterns.",
        "Avoid compulsive monitoring, moral scoring, or using the log as proof that a client should tolerate an unsafe condition.",
    ],
    [
        "Fawn-pattern inquiry",
        "Explore appeasement, over-functioning, conflict avoidance, and self-abandonment as context-dependent adaptive behavior.",
        "Fawn is descriptive framework language, not a DSM diagnosis or a universal female response. Assess power, culture, safety, and consequences before boundary experiments.",
    ],
]

PACING_MATRIX = [
    ["Signal", "Meaning for pacing", "Next move"],
    ["Client remains oriented and curious", "Current dose appears workable.", "Continue briefly, then integrate before adding intensity."],
    ["Speech speeds up, narrows, or becomes hard to follow", "Activation may be outrunning reflective capacity.", "Slow the pace; reduce questions; orient to immediate surroundings."],
    ["Blankness, collapse, distant gaze, or loss of words", "The client may have reduced access to choice or present-moment processing.", "Stop content exploration; use neutral orientation; assess safety and whether to end."],
    ["Breathing practice causes dizziness or air hunger", "The exercise is not currently tolerable or may be medically inappropriate.", "Return to normal breathing. Do not coach through symptoms. Refer for medical guidance when indicated."],
    ["Neck or body movement causes pain, tingling, weakness, or vertigo", "Movement is contraindicated until appropriately assessed.", "Stop movement. Use a non-movement option and recommend relevant medical evaluation."],
    ["Client repeatedly requests reassurance or permission", "The process may be shifting authority away from the client.", "Return choice, name options, and ask what information would support an autonomous decision."],
]

ADVERSE_RESPONSE = [
    ["Step", "Action"],
    ["1. Stop", "End the exercise or inquiry immediately. Do not frame persistence as courage or progress."],
    ["2. Orient", "Invite neutral observation of the room, date, location, stable surfaces, and ordinary sensory information. Do not force eye contact or breath control."],
    ["3. Assess", "Determine whether the client can communicate, choose, maintain safety, and return toward ordinary functioning."],
    ["4. Stabilize", "Reduce stimulation, offer water or a neutral pause when appropriate, and return to the client's established supports."],
    ["5. Escalate", "Use emergency, safeguarding, clinical consultation, or referral pathways when symptoms exceed scope or safety cannot be maintained."],
    ["6. Document", "Record the observable response, action taken, consultation or referral, and follow-up plan without speculative diagnosis."],
    ["7. Review", "Before future use, reconsider fit, dose, informed consent, practitioner competence, and whether the protocol should be discontinued."],
]

SYSTEMIC_REVIEW = [
    ["Dimension", "Practitioner inquiry", "Avoid"],
    ["Race and culture", "How do racialized expectations, surveillance, code-switching, family norms, or community belonging shape the pattern?", "Calling protective adaptation pathology or insisting on disclosure that increases risk."],
    ["Gender and identity", "How do gendered labor, safety calculations, identity threat, and social penalties affect available choices?", "Treating women as biologically uniform or presenting appeasement as an innate female trait."],
    ["Disability and neurodivergence", "Could sensory load, pain, fatigue, communication style, executive function, or access needs explain the response?", "Interpreting disability-related needs as resistance, avoidance, or dysregulation."],
    ["Economic and workplace power", "What material consequences limit the client's ability to remove or confront a stressor?", "Prescribing boundaries without considering income, housing, insurance, caregiving, or retaliation."],
    ["Faith and community", "What beliefs, practices, and communities support regulation, meaning, and belonging?", "Replacing a client's worldview with the practitioner's framework language."],
]

OUTCOME_REVIEW = [
    ["Outcome domain", "Useful evidence", "Do not infer"],
    ["Recovery", "Shorter return time after activation; more reliable use of supports.", "That the client is cured, trauma-free, or physiologically optimized."],
    ["Choice", "More ability to pause, decline, request time, or choose a response.", "That every difficult emotion should be regulated away."],
    ["Function", "Improved sleep routine, concentration, communication, follow-through, or daily stability as reported by the client.", "Medical improvement without appropriate measurement and clinical evaluation."],
    ["Environment", "Documented removal, restructuring, or containment of a repeat stressor.", "That internal practice can compensate for an unsafe system indefinitely."],
    ["Self-trust", "The client uses her own observations to make proportionate decisions.", "That agreement with the practitioner proves progress."],
]

REFERENCES = [
    "Substance Abuse and Mental Health Services Administration. Trauma-Informed Care in Behavioral Health Services (TIP 57). Updated 2025. https://www.samhsa.gov/resource/dbhis/tip-57-trauma-informed-care-behavioral-health-services",
    "Substance Abuse and Mental Health Services Administration. Trauma-Informed Approaches and Programs. Updated 2026. https://www.samhsa.gov/mental-health/trauma-violence/trauma-informed-approaches-programs",
    "National Institute for Health and Care Excellence. Post-traumatic stress disorder: NG116. Published 2018; reviewed 2025. https://www.nice.org.uk/guidance/ng116",
    "Balban, M. Y., et al. (2023). Brief structured respiration practices enhance mood and reduce physiological arousal. Cell Reports Medicine, 4(1), 100895. https://pubmed.ncbi.nlm.nih.gov/36630953/",
    "Laborde, S., Mosley, E., & Thayer, J. F. (2017). Heart Rate Variability and Cardiac Vagal Tone in Psychophysiological Research. Frontiers in Psychology, 8, 213. https://pubmed.ncbi.nlm.nih.gov/28265249/",
    "Quigley, K. S., et al. (2021). Functions of Interoception: From Energy Regulation to Experience of the Self. Trends in Neurosciences, 44(1), 29-38. https://pubmed.ncbi.nlm.nih.gov/33378654/",
    "Kuhfuss, M., et al. (2021). Somatic experiencing: effectiveness and key factors of a body-oriented trauma therapy. European Journal of Psychotraumatology, 12(1), 1929023. https://pubmed.ncbi.nlm.nih.gov/34290845/",
    "Manzotti, A., et al. (2024). An in-depth analysis of the polyvagal theory in light of current findings in neuroscience and clinical research. Developmental Psychobiology, 66(2), e22450. https://pubmed.ncbi.nlm.nih.gov/38388187/",
]


def rgb(value):
    return RGBColor.from_string(value)


def set_run(run, font="Aptos", size=10.5, color=BLACK, bold=False, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.width = Inches(widths[index] / 1440)


def paragraph(doc, text="", size=10.5, color=BLACK, bold=False, italic=False, align=None, before=0, after=6, line=1.2, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 11.5}
    before = {1: 8, 2: 8, 3: 6}
    p = paragraph(doc, text, size=sizes[level], color=DARK_PURPLE, bold=True, before=before[level], after=7, keep=True)
    p.style = f"Heading {level}"
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p


def label_value(doc, label, value):
    p = paragraph(doc, after=4)
    set_run(p.add_run(f"{label}: "), size=10.5, color=DARK_PURPLE, bold=True)
    set_run(p.add_run(value), size=10.5, color=BLACK)
    return p


def callout(doc, label, text, risk=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9936])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FDFBF6")
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(label.upper()), size=9, color=RISK if risk else GOLD, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    set_run(p2.add_run(text), size=10.5, color=BLACK)
    paragraph(doc, after=2)


def add_table(doc, rows, widths, font_size=8.8):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, DARK_PURPLE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, WARM_FILL)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            set_run(
                p.add_run(value),
                size=font_size,
                color=WHITE if row_index == 0 else BLACK,
                bold=row_index == 0 or (row_index > 0 and col_index == 0),
            )
    paragraph(doc, after=2)
    return table


def form_lines(doc, prompts, line_count=2):
    for prompt in prompts:
        paragraph(doc, prompt, size=10.5, color=DARK_PURPLE, bold=True, before=4, after=3, keep=True)
        for _ in range(line_count):
            p = paragraph(doc, "", after=7)
            p_pr = p._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "2")
            bottom.set(qn("w:color"), RULE)
            borders.append(bottom)
            p_pr.append(borders)


def page_break(doc):
    doc.add_page_break()


def add_page_number(paragraph_obj):
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph_obj.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run(run, size=8, color=DARK_PURPLE)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)

    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11.5)):
        style = doc.styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(DARK_PURPLE)

    list_style = doc.styles["List Bullet"]
    list_style.font.name = "Aptos"
    list_style.font.size = Pt(10.5)
    list_style.font.color.rgb = rgb(BLACK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(header.add_run("DISTINCT CHARACTER | PRACTITIONER LAYER"), size=8, color=DARK_PURPLE, bold=True)
    p_pr = header._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), GOLD)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer.paragraphs[0]
    set_run(footer.add_run("A. Solenne Institute | Confidential practitioner resource | "), size=8, color=DARK_PURPLE)
    add_page_number(footer)


def section_open(doc, kicker, title, intro):
    paragraph(doc, kicker.upper(), size=9, color=GOLD, bold=True, after=5)
    heading(doc, title, 1)
    paragraph(doc, intro, size=11, color=SLATE, after=10, line=1.25)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    if LOGO_PATH.exists():
        p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(0.82))
    paragraph(doc, "DISTINCT CHARACTER", size=11, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    title_p = paragraph(doc, TITLE, size=27, color=DARK_PURPLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=9, line=1.0)
    for run in title_p.runs:
        set_run(run, font="Georgia", size=27, color=DARK_PURPLE, bold=True)
    paragraph(doc, "Professional implementation, safety, and fidelity guidance", size=13, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    paragraph(doc, "PRACTITIONER-ONLY RESOURCE", size=10, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    paragraph(doc, f"Asset ID: {ASSET_ID} | Version 1.0 | 2026", size=9.5, color=DARK_PURPLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    callout(doc, "Restricted use", "This addendum is available only to approved Distinct Character practitioner-layer users and administrators. It does not grant a professional license, clinical credential, certification, or right to teach, reproduce, adapt, or sublicense the protocol.")
    paragraph(doc, "The Sovereign Bureau | A. Solenne Institute | A division of Granite Field Holdings Ltd. Co.", size=8.5, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, before=30, after=4)
    paragraph(doc, "Copyright 2026 Distinct Character. All rights reserved.", size=8.5, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    section_open(doc, "Rights and scope", "Professional Use Notice", "This document governs practitioner-supported use of the Somatic Baseline Protocol. It adds implementation boundaries and decision standards; it does not replace the client protocol, professional ethics, supervision, licensure rules, organizational policy, or applicable law.")
    heading(doc, "Purpose", 2)
    paragraph(doc, "The practitioner layer exists to help an appropriately qualified professional preserve client choice, monitor capacity, recognize limits, and translate protocol observations into proportionate next steps. The central standard is not intensity. It is accurate, ethical, and useful engagement.")
    heading(doc, "Non-treatment boundary", 2)
    paragraph(doc, "Distinct Character protocols are proprietary educational and behavioral-governance resources. They are not standalone medical or mental-health treatment, diagnostic instruments, crisis services, or substitutes for individualized care. Licensed professionals may integrate selected material only within their existing competence and scope.")
    heading(doc, "Rights boundary", 2)
    paragraph(doc, "No portion may be copied, distributed, taught, recorded, uploaded to another system, included in a client product, adapted for commercial use, or used to train an automated system without written authorization from A. Solenne Institute.")
    callout(doc, "Immediate safety", "If a client presents with imminent danger, inability to maintain safety, medical emergency, severe disorganization, or a condition requiring mandated action, stop protocol work and follow the emergency, safeguarding, reporting, consultation, and referral procedures required by your role and jurisdiction.", risk=True)
    page_break(doc)

    section_open(doc, "Document map", "Contents", "The addendum is organized around the decisions a practitioner must make before, during, and after using Somatic Baseline material with another person.")
    for item in CONTENTS:
        bullet(doc, item)
    heading(doc, "How to use this addendum", 2)
    paragraph(doc, "Read the scope, readiness, and adverse-response sections before practitioner-supported use. Use the implementation pages during planning and the printable tools during documentation, supervision, and completion review.")
    callout(doc, "Framework language", "Terms such as somatic stability, governance, zones, and the SDI belong to the Distinct Character framework. Use them as practical organizing language. Do not present them as established diagnoses, biomarkers, or universally accepted neurophysiological categories.")
    page_break(doc)

    section_open(doc, "Part I", "Professional Scope and Qualification Boundary", "The right use of the material depends on the practitioner's actual role, competence, setting, and legal authority. Practitioner-layer access does not broaden any of those boundaries.")
    add_table(doc, SCOPE_MATRIX, [2200, 3868, 3868], font_size=8.3)
    heading(doc, "Minimum operating standards", 2)
    for item in [
        "Use the protocol only with informed, voluntary participation and a clear explanation of your role.",
        "Remain within competence. Seek supervision or consultation when uncertainty affects safety or quality.",
        "Maintain confidentiality, secure records, and role-appropriate consent and documentation.",
        "Use referral pathways that are current, accessible, culturally responsive, and appropriate to the client's location.",
        "Do not claim guaranteed outcomes, biological correction, trauma resolution, or clinical efficacy from protocol completion.",
        "Do not use dependency, urgency, fear, or practitioner authority to pressure continuation.",
    ]:
        bullet(doc, item)
    page_break(doc)

    section_open(doc, "Part II", "Readiness, Consent, and Risk Classification", "A client may be interested in the work and still not be ready for a specific exercise, intensity, or practitioner relationship. Readiness is assessed repeatedly, not once.")
    add_table(doc, READINESS_GATE, [1900, 4018, 4018], font_size=8.1)
    page_break(doc)

    section_open(doc, "Risk decisions", "Green, Amber, and Red Response Levels", "Use the simplest classification that supports timely action. These levels organize practitioner response; they do not diagnose a condition.")
    add_table(doc, RISK_LEVELS, [1900, 3920, 4116], font_size=8.4)
    heading(doc, "Informed participation checklist", 2)
    for item in [
        "The client understands the educational purpose and the practitioner's role.",
        "The client knows that any prompt or exercise may be declined or stopped without penalty.",
        "The client understands confidentiality and its limits in the practitioner's setting.",
        "The client knows what to do if delayed distress occurs after the session.",
        "The practitioner has a role-appropriate plan for consultation, referral, safeguarding, and emergencies.",
        "Remote work includes a reasonable plan for location-specific support when the practitioner's role requires it.",
    ]:
        bullet(doc, item)
    callout(doc, "Touch and physical direction", "No touch is required to use this protocol. Practitioners must not introduce touch, manual manipulation, restraint, or physical treatment unless it is independently permitted by their profession, clinically appropriate, explicitly consented to, and governed by applicable standards.")
    page_break(doc)

    section_open(doc, "Part III", "Practitioner Implementation Architecture", "The session structure keeps the work narrow enough to monitor and concrete enough to integrate. One useful observation and one governed next step are preferable to an intense but poorly contained experience.")
    add_table(doc, SESSION_ARCHITECTURE, [1250, 4760, 3926], font_size=8.5)
    heading(doc, "The observation standard", 2)
    paragraph(doc, "Document what can be observed or what the client reports: changes in speech, attention, movement, breathing comfort, orientation, choice, distress, recovery time, and daily function. Separate observation from interpretation. A behavior may have multiple causes, including fatigue, pain, medication, disability, culture, environment, neurodivergence, grief, illness, or ordinary preference.")
    heading(doc, "The authority standard", 2)
    paragraph(doc, "The protocol should increase the client's capacity to notice, choose, and act. When the practitioner becomes the final authority on what the client's body means, the method has moved away from self-mastery. Offer hypotheses lightly, invite correction, and return decisions to the client whenever safety allows.")
    page_break(doc)

    section_open(doc, "Part IV", "Section-by-Section Clinical Adjacency Guidance", "The client protocol uses accessible shorthand to support learning. The practitioner is responsible for preserving uncertainty, avoiding overclaiming, and recognizing when another explanation or service is more appropriate.")
    add_table(doc, SECTION_GUIDANCE, [1950, 3890, 4096], font_size=7.9)
    page_break(doc)

    section_open(doc, "Evidence calibration", "What the Practitioner Must Not Overstate", "Practical language is useful only when it remains honest about the limits of current evidence and the limits of what can be inferred from an individual response.")
    for label, text in [
        ("Breathing", "Slow or exhale-focused breathing may help some people reduce arousal, but responses vary. It should not be described as guaranteed vagal activation, a universal treatment, or proof of nervous-system repair."),
        ("Heart-rate variability", "HRV is influenced by respiration, posture, fitness, age, medication, illness, measurement conditions, and other factors. Consumer wearables and single readings are not diagnostic."),
        ("Polyvagal theory", "The framework is influential in clinical practice and continues to be debated in anatomy, measurement, and interpretation. Use it as a model, not settled proof of why a client's behavior occurred."),
        ("Interoception", "Internal sensations can provide useful information, but they are not automatically accurate explanations of external events. Combine body data with context, evidence, health information, and reflective judgment."),
        ("Somatic practices", "Evidence differs across practices, populations, and outcomes. A client's immediate felt shift does not establish long-term efficacy, mechanism, or clinical treatment effect."),
        ("Identity and fawn patterns", "Avoid reducing complex behavior to biology or gender. Analyze history, culture, incentives, coercion, relationships, access, and material consequences."),
    ]:
        heading(doc, label, 3)
        paragraph(doc, text)
    page_break(doc)

    section_open(doc, "Part V", "Pacing, Titration, and Pendulation", "Titration and pendulation are pacing concepts, not proof that activation is therapeutic. The working question is whether the client retains enough orientation, choice, and recovery capacity for the current dose.")
    add_table(doc, PACING_MATRIX, [2250, 3330, 4356], font_size=8.2)
    heading(doc, "Low-capacity adaptation", 2)
    paragraph(doc, "When capacity is reduced, shorten the session target, use present-focused language, remove performance expectations, and prioritize ordinary functioning. A complete low-capacity session may consist of naming one signal, stopping one source of load, and choosing one support.")
    callout(doc, "Do not push through", "Panic, significant breathlessness, dizziness, pain, numbness, escalating unreality, confusion, or loss of choice are reasons to stop and assess. They are not evidence that the exercise is reaching a deeper layer.", risk=True)
    page_break(doc)

    section_open(doc, "Adverse response", "Stop, Orient, Assess, and Escalate", "A clear response pathway protects the client and the practitioner from improvising under pressure.")
    add_table(doc, ADVERSE_RESPONSE, [1250, 8686], font_size=8.8)
    heading(doc, "After-session contact", 2)
    paragraph(doc, "Do not promise availability beyond the boundaries of your service. Provide role-appropriate instructions before the session ends. If delayed symptoms emerge, determine whether the response belongs in routine follow-up, consultation, clinical referral, urgent assessment, or emergency response.")
    page_break(doc)

    section_open(doc, "Part VI", "Relational, Cultural, and Systemic Considerations", "Somatic information occurs inside power, culture, history, disability, economics, and relationships. Ethical use requires more than helping a person tolerate what harms her.")
    add_table(doc, SYSTEMIC_REVIEW, [1780, 4078, 4078], font_size=8.1)
    heading(doc, "Boundary experimentation", 2)
    paragraph(doc, "Begin with contexts where the client has meaningful choice and manageable consequences. Before encouraging direct confrontation, disclosure, reduced appeasement, or withdrawal of labor, review likely retaliation, dependence, employment, housing, immigration, custody, caregiving, health, and community consequences.")
    heading(doc, "Regulation is not compliance", 2)
    paragraph(doc, "A calmer presentation is not automatically a better outcome. The client may need anger, grief, distance, protest, legal information, advocacy, rest, medical care, or material change. Regulation should expand response options, not make a person easier for a harmful system to manage.")
    page_break(doc)

    section_open(doc, "Part VII", "Documentation, Outcome Review, and Fidelity", "Documentation should show what occurred, what decision was made, and why the next step was proportionate. It should not convert framework language into unsupported clinical fact.")
    add_table(doc, OUTCOME_REVIEW, [1800, 4000, 4136], font_size=8.2)
    heading(doc, "Documentation rules", 2)
    for item in [
        "Use the client's own words where meaning is important.",
        "Record observable signs and reported experience separately from practitioner interpretation.",
        "Document consent, pauses, declined exercises, adverse responses, consultation, referrals, and safety actions.",
        "Avoid writing that the SDI diagnosed dysregulation or that a zone label proved a biological state.",
        "Store records according to professional, contractual, privacy, retention, and breach-response requirements.",
        "Use only the minimum necessary personal information inside the portal's practitioner note fields.",
    ]:
        bullet(doc, item)
    page_break(doc)

    section_open(doc, "Printable tool 1", "Practitioner Readiness and Scope Review", "Complete before beginning practitioner-supported use and revisit whenever client needs, risk, setting, or practitioner role changes.")
    form_lines(doc, [
        "Practitioner name, role, credential if applicable, and jurisdiction",
        "Client identifier and protocol access confirmed",
        "Purpose of practitioner-supported use",
        "Needs that fall within my competence and role",
        "Known medical, mental-health, accessibility, communication, or environmental considerations relevant to pacing",
        "Current referral, consultation, safeguarding, and emergency pathways",
        "Consent and confidentiality discussion completed on",
    ], line_count=1)
    page_break(doc)

    section_open(doc, "Printable tool 2", "Session Preparation Record", "Use this page to define the smallest useful scope before the session begins.")
    form_lines(doc, [
        "Session date, location or delivery format, and planned duration",
        "Client-defined objective for this session",
        "Baseline activation or capacity rating and client language",
        "Protocol section or tool selected",
        "Why this selection is proportionate today",
        "Stop signals and adaptations agreed in advance",
        "Ordinary-life integration target",
    ], line_count=2)
    page_break(doc)

    section_open(doc, "Printable tool 3", "Practitioner Observation Record", "Record what was observed and reported without converting the observation into a diagnosis.")
    add_table(doc, [
        ["Observation point", "Before", "During", "After"],
        ["Orientation and attention", "", "", ""],
        ["Speech and communication", "", "", ""],
        ["Movement and posture", "", "", ""],
        ["Breathing comfort", "", "", ""],
        ["Choice and consent", "", "", ""],
        ["Reported distress or comfort", "", "", ""],
        ["Recovery and daily function", "", "", ""],
    ], [2200, 2578, 2578, 2580], font_size=8.3)
    form_lines(doc, ["Client's own words about the experience", "Practitioner interpretation, clearly labeled as a working hypothesis", "Action, adaptation, consultation, or referral"], line_count=2)
    page_break(doc)

    section_open(doc, "Printable tool 4", "Adverse Response and Referral Record", "Complete when an exercise is stopped for distress, loss of orientation, physical symptoms, risk, or a scope concern.")
    form_lines(doc, [
        "Date, time, setting, and activity underway",
        "Observable response and client's report",
        "Exercise stopped at",
        "Orientation or stabilization actions used",
        "Safety assessment and outcome",
        "Consultation, referral, safeguarding, or emergency action",
        "Follow-up responsibility, timing, and communication",
        "Changes required before any future protocol use",
    ], line_count=2)
    page_break(doc)

    section_open(doc, "Printable tool 5", "Supervision and Consultation Record", "Use consultation to improve decisions, not merely to obtain reassurance after the fact.")
    form_lines(doc, [
        "Question requiring consultation",
        "Relevant facts, client preferences, risk, and context",
        "What is inside and outside my role",
        "Consultant or supervisor and date",
        "Options considered",
        "Decision and rationale",
        "Client communication and follow-up",
    ], line_count=2)
    page_break(doc)

    section_open(doc, "Printable tool 6", "Fidelity Review", "Use at midpoint, completion, and whenever the process feels stuck or practitioner-led rather than client-governed.")
    for item in [
        "The client retains meaningful choice and can decline or stop without penalty.",
        "Protocol language is presented as framework language rather than diagnosis or settled mechanism.",
        "The selected exercise is inside my competence and proportionate to current capacity.",
        "I am not using regulation to encourage tolerance of an unsafe or exploitative condition.",
        "Client observations and practitioner interpretations remain distinct in documentation.",
        "Progress is evaluated through function, choice, recovery, and structural change rather than compliance with me.",
        "Referral, consultation, and escalation have occurred when needs exceeded scope.",
        "The client can describe how she will continue without dependence on the practitioner.",
    ]:
        label_value(doc, "Review", item)
        paragraph(doc, "Status:  Clear / Needs attention     Evidence or correction: ____________________________________________", size=9.5, after=8)
    page_break(doc)

    section_open(doc, "Part VIII", "Completion and Transition", "Practitioner-supported use is complete when the client can recognize relevant signals, choose proportionate actions, and identify when another kind of support is needed.")
    heading(doc, "Completion review", 2)
    for item in [
        "The client can distinguish observation from interpretation.",
        "The client can identify at least two usable supports and the conditions under which each should be used.",
        "The client can name stop signals and escalation criteria.",
        "At least one source of environmental friction has been removed, restructured, or addressed through an appropriate external pathway.",
        "The client has a realistic continuation plan that does not require constant practitioner access.",
        "Outstanding clinical, medical, legal, safety, or material needs have been referred or clearly documented.",
    ]:
        bullet(doc, item)
    heading(doc, "Transition note", 2)
    paragraph(doc, "Completion of Somatic Baseline does not require the absence of stress or difficult emotion. It requires a more accurate relationship to capacity, a safer method of responding to state changes, and clearer decisions about which problems need internal practice and which require structural or professional intervention.")
    callout(doc, "Final fidelity standard", "The practitioner succeeds when the client leaves with more authority, more accurate choices, and less dependence on the practitioner's interpretation.")
    page_break(doc)

    section_open(doc, "Evidence and context", "Selected References", "These sources support trauma-informed principles, evidence calibration, and professional caution. They do not validate the Distinct Character framework or the SDI as clinical instruments.")
    for index, reference in enumerate(REFERENCES, 1):
        p = paragraph(doc, after=6, line=1.15)
        set_run(p.add_run(f"{index}. "), size=9, color=DARK_PURPLE, bold=True)
        set_run(p.add_run(reference), size=9, color=BLACK)
    heading(doc, "Version history", 2)
    add_table(doc, [["Version", "Date", "Change"], ["1.0", "August 2026", "Initial practitioner-only therapeutic addendum created and activated for the Distinct Character portal."]], [1300, 1900, 6736], font_size=8.8)
    paragraph(doc, f"{TITLE} | {ASSET_ID}", size=9, color=DARK_PURPLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=14)

    doc.save(DOCX_PATH)
    source_copy = SOURCE_DIR / DOCX_PATH.name
    source_copy.write_bytes(DOCX_PATH.read_bytes())
    print(DOCX_PATH)
    print(source_copy)


if __name__ == "__main__":
    build()

