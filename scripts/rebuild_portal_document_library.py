from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import argparse
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
SOURCE_ROOT = ROOT / "source-library"
OUT_ROOT = ROOT / "output" / "portal-document-library-rebuild-2026-08-26"
DOCX_OUT = OUT_ROOT / "docx"
PDF_OUT = OUT_ROOT / "pdf"

SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")

DARK_PURPLE = RGBColor(45, 26, 75)
GOLD = RGBColor(184, 139, 54)
BLACK = RGBColor(18, 18, 18)
SLATE = RGBColor(56, 63, 72)
WHITE = RGBColor(255, 255, 255)

DARK_PURPLE_HEXES = {"2D1A4B", "2E1A47", "321D52", "3D2462", "3F2463", "40255F"}
LINE_PATTERN = re.compile(r"[_\.\-]{8,}")


@dataclass(frozen=True)
class DocumentSpec:
    pdf_name: str
    source: Path
    title: str
    special: str = ""


DOCUMENTS = [
    DocumentSpec("30-day-sovereignty-reset-protocol.pdf", SOURCE_ROOT / "30-day-sovereignty-reset" / "30_Day_Sovereignty_Reset_Protocol_Aligned_Revision_v0.1.docx", "30-Day Sovereignty Reset"),
    DocumentSpec("authority-framework-protocol.pdf", SOURCE_ROOT / "authority-framework" / "The_Authority_Framework_Protocol_Aligned_Revision_v0.1.docx", "Authority Framework"),
    DocumentSpec("biological-infrastructure-companion.pdf", SOURCE_ROOT / "foundation" / "DC_Biological_Infrastructure_Companion_Cover_Aligned_Revision_v1.2.docx", "Biological Infrastructure Companion"),
    DocumentSpec("body-signal-index.pdf", SOURCE_ROOT / "foundation" / "DC_Body_Signal_Index_Cover_Aligned_Revision_v0.1.docx", "Body Signal Index"),
    DocumentSpec("distinct-character-framework-glossary.pdf", SOURCE_ROOT / "foundation" / "Distinct_Character_Framework_Glossary_Cover_Aligned_Revision_v0.1.docx", "Distinct Character Framework Glossary"),
    DocumentSpec("enterprise-ip-mastermind-advisor-guide.pdf", SOURCE_ROOT / "enterprise-ip-mastermind" / "Enterprise_IP_Mastermind_Advisor_Legal_Ops_Guide_Cover_Aligned_Revision_v0.1.docx", "Enterprise IP Mastermind Advisor Guide"),
    DocumentSpec("enterprise-ip-mastermind-resource-suite.pdf", SOURCE_ROOT / "enterprise-ip-mastermind" / "Enterprise_IP_Mastermind_Resource_Suite_Cover_Aligned_Revision_v0.1.docx", "Enterprise IP Mastermind Resource Suite"),
    DocumentSpec("execution-architecture-companion.pdf", SOURCE_ROOT / "execution-architecture" / "EAP_Companion_Materials_Cover_Aligned_Revision_v0.1.docx", "Execution Architecture Companion"),
    DocumentSpec("execution-architecture-protocol.pdf", SOURCE_ROOT / "execution-architecture" / "Execution_Architecture_Protocol_Aligned_Revision_v0.1.docx", "Execution Architecture Protocol"),
    DocumentSpec("internal-signal-calibration-protocol.pdf", SOURCE_ROOT / "internal-signal-calibration" / "Internal_Signal_Calibration_Protocol_Aligned_Revision_v0.1.docx", "Internal Signal Calibration", "internal"),
    DocumentSpec("ios1-companion.pdf", SOURCE_ROOT / "identity-operating-system" / "IOS1_Companion_Materials_Cover_Aligned_Revision_v0.4.docx", "IOS-1 Companion"),
    DocumentSpec("ios1-protocol.pdf", SOURCE_ROOT / "identity-operating-system" / "IOS1_Identity_Operating_System_Protocol_Aligned_Revision_v0.1.docx", "Identity Operating System"),
    DocumentSpec("mes1-companion.pdf", SOURCE_ROOT / "masking-economy-system" / "MES1_Companion_Materials_Cover_Aligned_Revision_v0.4.docx", "MES-1 Companion"),
    DocumentSpec("mes1-protocol.pdf", SOURCE_ROOT / "masking-economy-system" / "MES1_Masking_Economy_System_Protocol_Aligned_Revision_v0.1.docx", "Masking Economy System", "mes"),
    DocumentSpec("ncs1-companion.pdf", SOURCE_ROOT / "narrative-control-system" / "NCS1_Companion_Materials_Cover_Aligned_Revision_v0.1.docx", "NCS-1 Companion", "ncs_companion"),
    DocumentSpec("ncs1-protocol.pdf", SOURCE_ROOT / "narrative-control-system" / "NCS-1_Narrative_Control_System_Aligned_Revision_v0.2.docx", "Narrative Control System", "ncs"),
    DocumentSpec("nervous-system-governance-guide.pdf", SOURCE_ROOT / "foundation" / "Nervous_System_Governance_Guide_Cover_Aligned_Revision_v0.1.docx", "Nervous System Governance Guide"),
    DocumentSpec("nsg-digestion-sleep-movement-recovery.pdf", SOURCE_ROOT / "foundation" / "NSG_Digestion_Sleep_Movement_Recovery_Cover_Aligned_Revision_v0.1.docx", "NSG Digestion, Sleep, Movement, Recovery"),
    DocumentSpec("self-mastery-blueprint-protocol.pdf", SOURCE_ROOT / "self-mastery-blueprint" / "Self_Mastery_Blueprint_Protocol_Aligned_Revision_v0.1.docx", "Self-Mastery Blueprint"),
    DocumentSpec("self-mastery-blueprint-relapse-reentry-ledger.pdf", SOURCE_ROOT / "self-mastery-blueprint" / "Self_Mastery_Blueprint_Relapse_ReEntry_Ledger_v0.1.docx", "Relapse and Re-Entry Ledger"),
    DocumentSpec("somatic-baseline-companion.pdf", SOURCE_ROOT / "somatic-baseline" / "SBP_Companion_Materials_Cover_Aligned_Revision_v0.4.docx", "Somatic Baseline Companion"),
    DocumentSpec("somatic-baseline-practitioner-therapeutic-addendum.pdf", SOURCE_ROOT / "somatic-baseline" / "Somatic_Baseline_Practitioner_Therapeutic_Addendum_v1.0.docx", "Somatic Baseline Practitioner Therapeutic Addendum"),
    DocumentSpec("somatic-baseline-protocol.pdf", SOURCE_ROOT / "somatic-baseline" / "Somatic_Baseline_Protocol_Aligned_Revision_v0.1.docx", "Somatic Baseline Protocol"),
]

# The 12 Dimensions file was intentionally designed as colored cards and its source is PDF.
PRESERVED_PDFS = [
    ("12-dimensions-wellness.pdf", SOURCE_ROOT / "foundation" / "12_Dimensions_Self_Mastery_Framework_v1.0.pdf"),
]

MES_TOC = [
    ("Orientation: How to Use This Protocol", "ORIENTATION"),
    ("Part I: The Classification", "PART I: THE CLASSIFICATION"),
    ("Part II: The Cost Structure", "PART II: THE COST STRUCTURE"),
    ("Part III: The Return Analysis", "PART III: THE RETURN ANALYSIS"),
    ("Part IV: Budget and Deployment", "PART IV: BUDGET AND DEPLOYMENT"),
    ("Part V: Protection and Governance", "PART V: PROTECTION AND GOVERNANCE"),
    ("Closing", "CLOSING"),
]

INTERNAL_TOC = [
    ("Section I: The Biological Reality of Intuition", "SECTION I"),
    ("Section II: Signal Distortion", "SECTION II"),
    ("Section III: Unblocking the Pathways", "SECTION III"),
    ("Section IV: Somatic Listening", "SECTION IV"),
    ("Section V: Tactical Execution", "SECTION V"),
    ("Section VI: Integration", "SECTION VI"),
    ("Bonus Tool: Weekly Somatic Intelligence Check-In", "BONUS TOOL"),
    ("Therapeutic Addendum", "THERAPEUTIC ADDENDUM"),
]

NCS_COMPANION_TOC = [
    ("Resource and Research Companion", "RESOURCE & RESEARCH COMPANION"),
    ("Integration Exercises", "INTEGRATION EXERCISES"),
    ("Quick Reference Guide", "QUICK REFERENCE GUIDE"),
    ("Optional Deepening Paths", "OPTIONAL DEEPENING PATHS"),
]

TRUE_BLANK_PAGES = {
    "biological-infrastructure-companion.pdf": [24],
    "execution-architecture-protocol.pdf": [12, 16],
    "internal-signal-calibration-protocol.pdf": [20],
    "nervous-system-governance-guide.pdf": [10],
    "nsg-digestion-sleep-movement-recovery.pdf": [43],
}


def compact(value: str) -> str:
    return " ".join(value.replace("\u00a0", " ").split())


def set_font(run, *, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Aptos")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def parse_hex(value: str | None) -> tuple[int, int, int] | None:
    if not value or value in {"auto", "none"} or len(value) != 6:
        return None
    try:
        return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        return None


def is_dark_hex(value: str | None) -> bool:
    rgb = parse_hex(value)
    if rgb is None:
        return False
    r, g, b = rgb
    luminance = (0.2126 * r + 0.7152 * g + 0.0722 * b) / 255
    return value.upper() in DARK_PURPLE_HEXES or luminance < 0.34


def needs_dark_text(rgb_value) -> bool:
    if rgb_value is None:
        return False
    try:
        r, g, b = (int(rgb_value[0]), int(rgb_value[1]), int(rgb_value[2]))
    except Exception:
        return False
    gold_like = r >= 130 and 80 <= g <= 185 and b <= 125
    pale_grey = max(r, g, b) - min(r, g, b) <= 16 and 105 <= r <= 235
    pale_colored = r >= 175 and g >= 155 and b >= 135
    return gold_like or pale_grey or pale_colored


def cell_fill(cell) -> str | None:
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def paragraph_fill(paragraph: Paragraph) -> str | None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=105, bottom=90, end=105) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph: Paragraph, color: str = "6A557F", size: int = 8, space: int = 2) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def insert_paragraph_before_table(table, text: str = "") -> Paragraph:
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    paragraph = Paragraph(p, table._parent)
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    new_paragraph = Paragraph(p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def remove_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def normalize_quotes(text: str) -> str:
    text = compact(text)
    if len(text) >= 2 and text[0] in {'"', "\u201c", "\u2018"} and text[-1] in {'"', "\u201d", "\u2019"}:
        return text[1:-1].strip()
    return text


def looks_like_prompt(text: str) -> bool:
    value = compact(text)
    if not value or len(value) > 420:
        return False
    prefixes = (
        "when ", "what ", "where ", "which ", "who ", "how ", "why ", "describe ",
        "identify ", "record ", "list ", "note ", "notes", "reflection", "scenario",
        "context description", "masks required", "current ", "complete ", "write ",
    )
    return "?" in value or value.lower().startswith(prefixes) or value.endswith(":")


def is_decorative_callout(table) -> bool:
    if len(table.rows) != 1 or len(table.rows[0].cells) != 2:
        return False
    cells = table.rows[0].cells
    fills = [cell_fill(cell) for cell in cells]
    text = compact(" ".join(cell.text for cell in cells))
    if not text or len(text) > 1300:
        return False
    has_dark = any(is_dark_hex(fill) for fill in fills)
    note_language = any(token in text.lower() for token in ("note", "prerequisite", "remember", "governance", "important"))
    return has_dark and (note_language or not compact(cells[0].text))


def is_prompt_box(table) -> bool:
    if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
        return False
    text = compact(table.rows[0].cells[0].text)
    if not looks_like_prompt(text):
        return False
    row_height = table.rows[0]._tr.get_or_add_trPr().find(qn("w:trHeight"))
    height = int(row_height.get(qn("w:val"), "0")) if row_height is not None else 0
    return height >= 420 or len(table.rows[0].cells[0].paragraphs) >= 2


def unbox_callout(table) -> None:
    text_parts = [normalize_quotes(cell.text) for cell in table.rows[0].cells if compact(cell.text)]
    text = "\n".join(part for part in text_parts if part)
    paragraph = insert_paragraph_before_table(table)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_font(run, size=9.5, color=DARK_PURPLE, bold=(index == 0), italic=False)
    remove_table(table)


def unbox_prompt(table) -> None:
    text = compact(table.rows[0].cells[0].text)
    prompt = insert_paragraph_before_table(table)
    prompt.paragraph_format.space_before = Pt(6)
    prompt.paragraph_format.space_after = Pt(2)
    run = prompt.add_run(text)
    set_font(run, size=9.5, color=DARK_PURPLE, bold=True, italic=True)

    row_height = table.rows[0]._tr.get_or_add_trPr().find(qn("w:trHeight"))
    height = int(row_height.get(qn("w:val"), "0")) if row_height is not None else 900
    line_count = max(3, min(7, round(height / 250)))
    anchor = prompt
    for _ in range(line_count):
        anchor = insert_paragraph_after(anchor)
        anchor.paragraph_format.space_after = Pt(7)
        anchor.paragraph_format.line_spacing = 1
        set_paragraph_bottom_border(anchor)
    remove_table(table)


def remove_leading_blank_section(doc: Document) -> int:
    removed = 0
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            break
        if child.tag != qn("w:p"):
            break
        text = "".join(child.itertext()).strip()
        sect_pr = child.find(f".//{qn('w:sectPr')}")
        if text or sect_pr is None:
            break
        body.remove(child)
        removed += 1
    return removed


def remove_duplicate_break_before(doc: Document, target_text: str) -> int:
    paragraphs = list(doc.paragraphs)
    target_index = next((i for i, p in enumerate(paragraphs) if compact(p.text).upper() == target_text.upper()), None)
    if target_index is None:
        return 0
    candidates = []
    for paragraph in reversed(paragraphs[max(0, target_index - 8) : target_index]):
        if compact(paragraph.text):
            break
        breaks = paragraph._p.findall(f".//{qn('w:br')}")
        if any(node.get(qn("w:type")) == "page" for node in breaks):
            candidates.append(paragraph)
    if len(candidates) < 2:
        return 0
    for paragraph in candidates[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    return len(candidates) - 1


def move_empty_page_breaks_to_next_paragraph(doc: Document) -> int:
    body = doc.element.body
    changed = 0
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.itertext()).strip()
        breaks = child.findall(f".//{qn('w:br')}")
        page_breaks = [node for node in breaks if node.get(qn("w:type")) == "page"]
        if text or not page_breaks:
            continue
        next_node = child.getnext()
        while next_node is not None and next_node.tag == qn("w:p") and not "".join(next_node.itertext()).strip():
            next_node = next_node.getnext()
        if next_node is not None and next_node.tag == qn("w:p"):
            p_pr = next_node.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                next_node.insert(0, p_pr)
            if p_pr.find(qn("w:pageBreakBefore")) is None:
                p_pr.append(OxmlElement("w:pageBreakBefore"))
        body.remove(child)
        changed += 1
    return changed


def remove_ncs_front_matter_spacer(doc: Document) -> int:
    paragraphs = list(doc.paragraphs)
    marker = next((i for i, paragraph in enumerate(paragraphs) if "LETTER FROM AZARI SOLENNE" in paragraph.text.upper()), None)
    if marker is None or marker == 0:
        return 0
    removed = 0
    for paragraph in reversed(paragraphs[max(0, marker - 3) : marker]):
        if compact(paragraph.text):
            continue
        p_pr = paragraph._p.pPr
        spacing = p_pr.find(qn("w:spacing")) if p_pr is not None else None
        before = int(spacing.get(qn("w:before"), "0")) if spacing is not None else 0
        after = int(spacing.get(qn("w:after"), "0")) if spacing is not None else 0
        if before + after >= 1200:
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1
    return removed


def collapse_empty_section_break_paragraphs(doc: Document) -> int:
    body = doc.element.body
    changed = 0
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        sect_pr = child.find(f".//{qn('w:sectPr')}")
        if sect_pr is None:
            continue
        text = compact("".join(child.itertext()))
        footer_like = (
            not text
            or text.lower().startswith("the sovereign bureau")
            or (text.lower().startswith("distinct character") and len(text) < 140)
        )
        if not footer_like:
            continue
        previous = child.getprevious()
        while previous is not None and previous.tag != qn("w:p"):
            previous = previous.getprevious()
        if previous is None:
            continue
        previous_p_pr = previous.find(qn("w:pPr"))
        if previous_p_pr is None:
            previous_p_pr = OxmlElement("w:pPr")
            previous.insert(0, previous_p_pr)
        old = previous_p_pr.find(qn("w:sectPr"))
        if old is not None:
            previous_p_pr.remove(old)
        previous_p_pr.append(deepcopy(sect_pr))
        body.remove(child)
        changed += 1
    return changed


def usable_width(doc: Document) -> int:
    section = doc.sections[-1]
    return int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)


def table_grid_widths(table) -> list[int]:
    grid = table._tbl.tblGrid
    widths = []
    if grid is not None:
        for col in grid.findall(qn("w:gridCol")):
            try:
                widths.append(int(col.get(qn("w:w"))))
            except (TypeError, ValueError):
                pass
    if widths:
        return widths
    if not table.rows:
        return []
    for cell in table.rows[0].cells:
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        try:
            widths.append(int(tc_w.get(qn("w:w"))))
        except (AttributeError, TypeError, ValueError):
            widths.append(1)
    return widths


def scale_table_preserving_proportions(table, target_width: int) -> None:
    widths = table_grid_widths(table)
    if not widths:
        return
    total = max(1, sum(widths))
    scaled = [max(360, round(width * target_width / total)) for width in widths]
    scaled[-1] += target_width - sum(scaled)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(target_width))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is not None:
        for col, width in zip(grid.findall(qn("w:gridCol")), scaled):
            col.set(qn("w:w"), str(width))

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        for height in tr_pr.findall(qn("w:trHeight")):
            tr_pr.remove(height)
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and len(table.rows) > 2 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)


def fix_table_contrast(table) -> int:
    changed = 0
    for row in table.rows:
        for cell in row.cells:
            dark = is_dark_hex(cell_fill(cell))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.widow_control = True
                for run in paragraph.runs:
                    if dark:
                        set_font(run, color=WHITE)
                        changed += 1
                    elif needs_dark_text(run.font.color.rgb):
                        set_font(run, color=DARK_PURPLE)
                        changed += 1
    return changed


def expand_writer_lines(paragraph: Paragraph) -> int:
    text = paragraph.text
    if not LINE_PATTERN.search(text):
        return 0
    label = LINE_PATTERN.sub("", text).strip()
    paragraph.clear()
    if label:
        run = paragraph.add_run(label)
        set_font(run, size=9.5, color=DARK_PURPLE, bold=True)
        line = insert_paragraph_after(paragraph)
    else:
        line = paragraph
    line.paragraph_format.space_after = Pt(7)
    set_paragraph_bottom_border(line)
    return 1


def normalize_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer, section.even_page_header, section.even_page_footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    set_font(run, color=DARK_PURPLE if part in (section.header, section.first_page_header, section.even_page_header) else BLACK)


def normalize_section_starts(doc: Document) -> int:
    changed = 0
    for section in doc.sections:
        if section.start_type in {WD_SECTION_START.ODD_PAGE, WD_SECTION_START.EVEN_PAGE}:
            section.start_type = WD_SECTION_START.NEW_PAGE
            changed += 1
    return changed


def normalize_body_text_contrast(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if needs_dark_text(run.font.color.rgb):
                set_font(run, color=DARK_PURPLE)
                changed += 1
    return changed


def normalize_shaded_paragraph_contrast(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        if not is_dark_hex(paragraph_fill(paragraph)):
            continue
        for run in paragraph.runs:
            if compact(run.text):
                set_font(run, color=WHITE)
                changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not is_dark_hex(paragraph_fill(paragraph)):
                        continue
                    for run in paragraph.runs:
                        if compact(run.text):
                            set_font(run, color=WHITE)
                            changed += 1
    return changed


def force_ncs_cover_text_white(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = compact(paragraph.text).upper()
        if text == "THE SOVEREIGN BUREAU":
            break
        for run in paragraph.runs:
            if compact(run.text):
                set_font(run, color=WHITE)
                changed += 1
    return changed


def clear_between(doc: Document, start: Paragraph, stop_text: str, *, stop_heading_only: bool = False) -> None:
    node = start._p.getnext()
    while node is not None:
        next_node = node.getnext()
        text = compact("".join(node.itertext()))
        p_pr = node.find(qn("w:pPr")) if node.tag == qn("w:p") else None
        p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
        is_heading = p_style is not None and str(p_style.get(qn("w:val"), "")).lower().startswith("heading")
        if text.upper().startswith(stop_text.upper()) and (not stop_heading_only or is_heading):
            p_pr = node.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                node.insert(0, p_pr)
            if p_pr.find(qn("w:pageBreakBefore")) is None:
                p_pr.append(OxmlElement("w:pageBreakBefore"))
            break
        doc.element.body.remove(node)
        node = next_node


def add_static_toc(
    doc: Document,
    *,
    heading_text: str,
    stop_text: str,
    entries: list[tuple[str, str]],
    page_numbers: dict[str, int] | None = None,
    stop_heading_only: bool = False,
) -> None:
    heading = next(paragraph for paragraph in doc.paragraphs if compact(paragraph.text).upper() == heading_text.upper())
    clear_between(doc, heading, stop_text, stop_heading_only=stop_heading_only)
    anchor = heading
    tab_position = Inches(6.25)
    for label, key in entries:
        paragraph = insert_paragraph_after(anchor)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        number = str(page_numbers.get(key, 0)) if page_numbers else "00"
        run = paragraph.add_run(f"{label}\t{number}")
        set_font(run, size=11, color=DARK_PURPLE, bold=True)
        anchor = paragraph


def add_special_toc(doc: Document, special: str, page_numbers: dict[str, int] | None = None) -> list[tuple[str, str]]:
    if special == "mes":
        add_static_toc(doc, heading_text="CONTENTS", stop_text="ORIENTATION", entries=MES_TOC, page_numbers=page_numbers)
        return MES_TOC
    if special == "internal":
        add_static_toc(doc, heading_text="CONTENTS", stop_text="WELCOME TO YOUR PROTOCOL", entries=INTERNAL_TOC, page_numbers=page_numbers)
        return INTERNAL_TOC
    if special == "ncs_companion":
        add_static_toc(doc, heading_text="CONTENTS", stop_text="RESOURCE & RESEARCH COMPANION", entries=NCS_COMPANION_TOC, page_numbers=page_numbers, stop_heading_only=True)
        return NCS_COMPANION_TOC
    return []


def locate_pdf_pages(pdf_path: Path, keys: list[str]) -> dict[str, int]:
    reader = PdfReader(str(pdf_path))
    pages = [compact(page.extract_text() or "").upper() for page in reader.pages]
    contents_index = next((index for index, text in enumerate(pages) if "CONTENTS" in text), -1)
    found: dict[str, int] = {}
    for key in keys:
        needle = compact(key).upper()
        for index, text in enumerate(pages, start=1):
            if index <= contents_index + 1:
                continue
            if needle in text:
                found[key] = index
                break
    return found


def convert_to_pdf(docx_path: Path, target_pdf: Path) -> None:
    target_pdf.parent.mkdir(parents=True, exist_ok=True)
    profile = Path(tempfile.mkdtemp(prefix="dc_library_rebuild_"))
    try:
        result = subprocess.run(
            [
                str(SOFFICE), "--headless", "--norestore", "--nodefault", "--nolockcheck",
                f"-env:UserInstallation={profile.as_uri()}", "--convert-to", "pdf", "--outdir",
                str(target_pdf.parent), str(docx_path),
            ],
            cwd=str(ROOT), capture_output=True, text=True, timeout=420,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice failed for {docx_path.name}: {result.stderr}\n{result.stdout}")
        generated = target_pdf.parent / f"{docx_path.stem}.pdf"
        if not generated.exists():
            raise FileNotFoundError(generated)
        if generated != target_pdf:
            if target_pdf.exists():
                target_pdf.unlink()
            generated.replace(target_pdf)
    finally:
        shutil.rmtree(profile, ignore_errors=True)


def strip_pdf_pages(pdf_path: Path, page_numbers: list[int]) -> None:
    if not page_numbers:
        return
    remove = {number - 1 for number in page_numbers}
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()
    for index, page in enumerate(reader.pages):
        if index not in remove:
            writer.add_page(page)
    temporary = pdf_path.with_suffix(".stripped.pdf")
    with temporary.open("wb") as handle:
        writer.write(handle)
    temporary.replace(pdf_path)


def adjust_for_removed_pages(page_numbers: dict[str, int], removed_pages: list[int]) -> dict[str, int]:
    adjusted = {}
    for key, page in page_numbers.items():
        adjusted[key] = page - sum(1 for removed in removed_pages if removed < page)
    return adjusted


def rebuild_document(spec: DocumentSpec) -> tuple[Path, Path, dict[str, int]]:
    if not spec.source.exists():
        raise FileNotFoundError(spec.source)
    doc = Document(spec.source)
    metrics = {"leading_blank_sections": 0, "duplicate_breaks": 0, "page_breaks_normalized": 0, "section_breaks_collapsed": 0, "section_starts_fixed": 0, "callouts_removed": 0, "prompt_boxes_removed": 0, "contrast_runs_fixed": 0, "writer_lines_fixed": 0, "tables_scaled": 0}

    metrics["section_breaks_collapsed"] = collapse_empty_section_break_paragraphs(doc)
    metrics["section_starts_fixed"] = normalize_section_starts(doc)
    metrics["page_breaks_normalized"] = move_empty_page_breaks_to_next_paragraph(doc)

    if spec.special == "ncs":
        metrics["leading_blank_sections"] = remove_leading_blank_section(doc)
        metrics["leading_blank_sections"] += remove_ncs_front_matter_spacer(doc)
    toc_entries = add_special_toc(doc, spec.special)
    if spec.special == "mes":
        metrics["duplicate_breaks"] = remove_duplicate_break_before(doc, "CONTENTS")

    for table in list(doc.tables):
        if is_prompt_box(table):
            unbox_prompt(table)
            metrics["prompt_boxes_removed"] += 1
            continue
        if is_decorative_callout(table):
            unbox_callout(table)
            metrics["callouts_removed"] += 1
            continue

    target_width = usable_width(doc)
    for table in doc.tables:
        scale_table_preserving_proportions(table, target_width)
        metrics["tables_scaled"] += 1
        metrics["contrast_runs_fixed"] += fix_table_contrast(table)

    for paragraph in list(doc.paragraphs):
        metrics["writer_lines_fixed"] += expand_writer_lines(paragraph)
        paragraph.paragraph_format.widow_control = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    metrics["writer_lines_fixed"] += expand_writer_lines(paragraph)

    metrics["contrast_runs_fixed"] += normalize_body_text_contrast(doc)
    metrics["contrast_runs_fixed"] += normalize_shaded_paragraph_contrast(doc)
    if spec.special == "ncs":
        metrics["contrast_runs_fixed"] += force_ncs_cover_text_white(doc)
    normalize_header_footer(doc)

    docx_target = DOCX_OUT / spec.source.name.replace(".docx", "_Portal_QA_Corrected.docx")
    pdf_target = PDF_OUT / spec.pdf_name
    docx_target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_target)
    convert_to_pdf(docx_target, pdf_target)

    if toc_entries:
        page_numbers = locate_pdf_pages(pdf_target, [key for _, key in toc_entries])
        page_numbers = adjust_for_removed_pages(page_numbers, TRUE_BLANK_PAGES.get(spec.pdf_name, []))
        doc = Document(docx_target)
        add_special_toc(doc, spec.special, page_numbers)
        doc.save(docx_target)
        convert_to_pdf(docx_target, pdf_target)

    strip_pdf_pages(pdf_target, TRUE_BLANK_PAGES.get(spec.pdf_name, []))

    return docx_target, pdf_target, metrics


def pdf_audit(pdf_path: Path) -> dict[str, object]:
    reader = PdfReader(str(pdf_path))
    texts = [compact(page.extract_text() or "") for page in reader.pages]
    blank_pages = [index for index, text in enumerate(texts, start=1) if len(text) < 18]
    blank_toc_pages = [index for index, text in enumerate(texts, start=1) if "CONTENTS" in text.upper() and len(text) < 100]
    return {
        "pages": len(reader.pages),
        "blank_pages": blank_pages,
        "blank_toc_pages": blank_toc_pages,
        "first_page_text": len(texts[0]) if texts else 0,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--only", help="Comma-separated PDF filenames to rebuild")
    args = parser.parse_args()
    selected = {item.strip() for item in args.only.split(",")} if args.only else None

    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    PDF_OUT.mkdir(parents=True, exist_ok=True)
    report = [
        "# Portal Document Library Rebuild and Audit",
        "",
        "Rules: preserve source hierarchy and column proportions; remove decorative quote/callout boxes; replace boxed writing areas with full-width ruled lines; enforce white text on dark purple; prevent table row splitting; keep headers and footers readable.",
        "",
        "## Rebuilt documents",
        "",
    ]

    for spec in DOCUMENTS:
        if selected is not None and spec.pdf_name not in selected:
            continue
        docx_path, pdf_path, metrics = rebuild_document(spec)
        audit = pdf_audit(pdf_path)
        report.append(
            f"- `{spec.pdf_name}`: {audit['pages']} pages; blank pages={audit['blank_pages']}; blank TOCs={audit['blank_toc_pages']}; "
            f"callouts removed={metrics['callouts_removed']}; prompt boxes removed={metrics['prompt_boxes_removed']}; "
            f"contrast fixes={metrics['contrast_runs_fixed']}; writer-line fixes={metrics['writer_lines_fixed']}; tables preserved/scaled={metrics['tables_scaled']}."
        )

    report.extend(["", "## Preserved card-based framework", ""])
    for pdf_name, source in PRESERVED_PDFS:
        if selected is not None and pdf_name not in selected:
            continue
        if not source.exists():
            raise FileNotFoundError(source)
        target = PDF_OUT / pdf_name
        shutil.copy2(source, target)
        audit = pdf_audit(target)
        report.append(f"- `{pdf_name}`: preserved branded card colors and structure; {audit['pages']} pages; blank pages={audit['blank_pages']}.")

    report_path = OUT_ROOT / "portal-document-library-audit.md"
    report_path.write_text("\n".join(report), encoding="utf-8")
    print(report_path)


if __name__ == "__main__":
    main()

