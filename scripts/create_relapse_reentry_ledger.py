from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    LongTable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "source-library" / "self-mastery-blueprint"
DOCX_PATH = OUT_DIR / "Self_Mastery_Blueprint_Relapse_ReEntry_Ledger_v0.1.docx"
PDF_PATH = OUT_DIR / "Self_Mastery_Blueprint_Relapse_ReEntry_Ledger_v0.1.pdf"
LOGO_PATH = ROOT / "public" / "assets" / "dc-logo.png"

DARK_PURPLE = "#21162E"
GOLD = "#BFA46A"
BLACK = "#111111"
SLATE = "#374151"
LIGHT_RULE = "#D9D2C1"
LIGHT_FILL = "#F7F5EF"
WHITE = "#FFFFFF"


CONTENT = {
    "title": "The Relapse & Re-Entry Ledger",
    "subtitle": "A Self-Mastery Blueprint tool for restoring behavioral governance after disruption.",
    "purpose": [
        "This ledger exists because self-mastery is not proven by never losing alignment. It is proven by how accurately the system returns after disruption.",
        "Inside the Self-Mastery Blueprint, relapse is not treated as moral failure, weak discipline, or evidence that progress was false. Relapse is treated as governance data. It shows where a standard, boundary, identity pattern, nervous system state, relational pressure, environmental condition, or execution structure could not hold under load.",
        "The work is not to punish the breach. The work is to classify what happened, restore command, protect the next decision, and update the architecture so the same collapse does not become the default path.",
    ],
    "use_rules": [
        "Use this ledger after a breach, collapse, avoidance cycle, emotional override, relational over-functioning pattern, or return to an old identity script.",
        "Do not use this ledger to self-attack. Shame reduces learning accuracy and often reinforces the same pattern.",
        "Record facts before interpretation. The first goal is cleaner data.",
        "Use the low-capacity version when the full ledger feels too large.",
        "Seek professional support when safety, trauma symptoms, addiction, abuse, or severe mental health distress are present.",
    ],
    "reentry_steps": [
        "Stabilize state. Eat, hydrate, sleep, breathe, orient, move gently, reduce sensory load, or step away from the triggering environment.",
        "Name the breach without drama. Write one factual sentence about what happened.",
        "Classify the system area. Choose one primary category from the Relapse Classification Map.",
        "Choose one re-entry action. It must be small, concrete, and possible within the next 24 hours.",
        "Update the architecture. Identify the one structure that would make the same relapse less likely next time.",
    ],
}

CORE_DISTINCTIONS = [
    ["Term", "Definition", "Governance Response"],
    ["Lapse", "A temporary interruption in the chosen standard or practice.", "Return quickly. Reduce shame. Record the condition."],
    ["Relapse", "A repeated or deeper return to an old operating pattern.", "Classify the system failure. Rebuild the support structure."],
    ["Regression", "A broader decline in capacity, clarity, or consistency across multiple areas.", "Stabilize biology first. Narrow the operating load."],
    ["Re-entry", "The deliberate process of returning to command after disruption.", "Use the smallest truthful action that restores direction."],
]

CLASSIFICATION_MAP = [
    ["System Area", "Relapse Signal", "What It May Indicate", "Re-Entry Question"],
    ["Biological Capacity", "Exhaustion, shutdown, activation, poor recovery, low frustration tolerance.", "The body was operating below the level required for the standard.", "What biological support was missing before the breach?"],
    ["Identity Standard", "Acting from an old role, self-abandonment, approval seeking, or familiar defeat.", "The current identity standard was not strong enough under pressure.", "Which version of self took command?"],
    ["Execution Architecture", "Avoidance, open loops, procrastination, urgency cycles, overbuilding.", "The action system had too much friction or too little clarity.", "What is the smallest executable next action?"],
    ["Relational Command", "Over-explaining, appeasing, rescuing, boundary leakage, resentment.", "Relational pressure overrode self-governance.", "What boundary or consequence was not protected?"],
    ["Narrative Control", "Catastrophizing, inherited meaning, identity-level conclusions.", "Interpretation moved faster than review.", "What story became truth without verification?"],
    ["Environment", "Repeated exposure to friction, chaos, access, temptation, or interruption.", "The environment made the old pattern easier than the new standard.", "What must be removed, restructured, or protected?"],
    ["Self-Efficacy", "A collapse in agency, learned helplessness, or distrust of follow-through.", "The system needs evidence of small successful action.", "What action can rebuild proof without pressure?"],
]

LOW_CAPACITY = [
    ["Prompt", "One-Line Entry"],
    ["What happened?", "One factual sentence."],
    ["What system area was affected?", "Biology, identity, execution, relational command, narrative, environment, or self-efficacy."],
    ["What is the smallest return-to-command action?", "One action that can be completed today."],
    ["What must be protected for the next 24 hours?", "One boundary, condition, or reduction."],
]

LEDGER_TEMPLATE = [
    ["Ledger Field", "Client Entry"],
    ["Date and Context", "When did the breach happen? What was occurring before it?"],
    ["State Before the Breach", "What was your body state, energy level, emotional state, and cognitive load?"],
    ["Observed Breach", "What action, avoidance, reaction, collapse, or old pattern occurred?"],
    ["System Area Affected", "Biological capacity, identity, execution, relational command, narrative control, environment, boundary, or self-efficacy."],
    ["Trigger or Friction", "What condition made the old pattern more available?"],
    ["Cost", "What did the breach cost in energy, clarity, time, trust, money, health, focus, or relational authority?"],
    ["Protected Function", "What was the old pattern trying to protect, avoid, preserve, or secure?"],
    ["Re-Entry Action", "What single action restores direction without forcing a full reset?"],
    ["Architecture Update", "What structure, boundary, practice, environment change, or decision rule must be adjusted?"],
    ["Evidence After Re-Entry", "What changed after the re-entry action? What evidence shows restored command?"],
]

REENTRY_PLAN = [
    ["Area", "Decision", "Boundary", "Completion Evidence"],
    ["Biology", "What will support enough capacity?", "What load must be reduced?", "What proves the body is more stable?"],
    ["Behavior", "What is the next governed action?", "What distraction or avoidance route is blocked?", "What proves action resumed?"],
    ["Relational", "What communication or silence is required?", "What over-functioning will not be funded?", "What proves command was protected?"],
    ["Environment", "What must be changed around you?", "What access, friction, or trigger must be limited?", "What proves the environment supports re-entry?"],
]

PATTERN_REVIEW = [
    ["Review Question", "What to Look For", "System Decision"],
    ["What relapse category repeats most often?", "The same system area appearing across entries.", "Strengthen that layer before adding more goals."],
    ["What state usually comes first?", "Fatigue, activation, shutdown, resentment, urgency, or numbness.", "Add earlier biological or relational intervention."],
    ["What old identity script returns?", "Good girl, rescuer, performer, invisible one, overachiever, abandoned self, or powerless self.", "Name the script and define the replacement standard."],
    ["What makes re-entry successful?", "Specific support, simpler action, reduced exposure, boundaries, or repair conversations.", "Formalize the condition as part of the command center."],
    ["What must no longer be optional?", "Sleep, food, boundary, review cycle, money standard, communication limit, or pacing rule.", "Convert it into a protected operating rule."],
]


def docx_font(run, size=None, color=None, bold=None):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bold is not None:
        run.bold = bold


def docx_paragraph(doc, text, size=10.5, color=BLACK, bold=False, align=None, before=0, after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    r = p.add_run(text)
    docx_font(r, size=size, color=color, bold=bold)
    return p


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def style_docx_table(table):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_fill(cell, DARK_PURPLE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    docx_font(run, size=8.5, color=WHITE if row_index == 0 else BLACK, bold=row_index == 0)


def add_docx_table(doc, rows):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_docx_table(table)


def create_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    header = section.header.paragraphs[0]
    header.text = "Distinct Character | Self-Mastery Blueprint"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    docx_font(header.runs[0], size=8.5, color=DARK_PURPLE)
    footer = section.footer.paragraphs[0]
    footer.text = "Copyright 2026 A. Solenne Institute. All rights reserved."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docx_font(footer.runs[0], size=8, color=SLATE)

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(0.85))
    docx_paragraph(doc, "DISTINCT CHARACTER", size=11, color=GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    docx_paragraph(doc, CONTENT["title"], size=25, color=DARK_PURPLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    docx_paragraph(doc, CONTENT["subtitle"], size=12, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    docx_paragraph(doc, "Self-Mastery Blueprint Companion Resource | Version 1.0 | 2026", size=9.5, color=DARK_PURPLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    docx_paragraph(doc, "Copyright and Use Notice", size=18, color=DARK_PURPLE, bold=True, after=8)
    for text in [
        "The Relapse & Re-Entry Ledger is a proprietary educational resource within the Distinct Character Self-Mastery Blueprint. It is intended for personal protocol use, practitioner-supported review, and approved Distinct Character implementation contexts.",
        "No portion of this document may be copied, distributed, modified, taught, licensed, sold, or repackaged without written permission from A. Solenne Institute.",
        "Copyright 2026 A. Solenne Institute. All rights reserved. Distinct Character and related protocol language are part of the Distinct Character ecosystem.",
    ]:
        docx_paragraph(doc, text)
    docx_paragraph(doc, "Educational and Medical Disclaimer", size=13, color=DARK_PURPLE, bold=True, before=10)
    for text in [
        "This document is educational and reflective in nature. It does not provide medical advice, mental health diagnosis, crisis care, or individualized treatment. It is not a substitute for care from a qualified medical, mental health, legal, or financial professional.",
        "If you are experiencing danger, self-harm thoughts, severe distress, medical symptoms, abuse, coercion, or inability to maintain safety, stop using this resource and seek immediate support from emergency services, crisis resources, or a qualified professional.",
        "Use low-capacity adaptations when the full process is too much. Re-entry should restore stability and function, not intensify pressure.",
    ]:
        docx_paragraph(doc, text)
    doc.add_page_break()

    docx_paragraph(doc, "Purpose", size=18, color=DARK_PURPLE, bold=True)
    for text in CONTENT["purpose"]:
        docx_paragraph(doc, text)
    docx_paragraph(doc, "Core Distinction", size=13, color=DARK_PURPLE, bold=True, before=8)
    add_docx_table(doc, CORE_DISTINCTIONS)
    docx_paragraph(doc, "Use Rules", size=13, color=DARK_PURPLE, bold=True, before=8)
    for item in CONTENT["use_rules"]:
        docx_paragraph(doc, item, style="List Bullet", after=3)
    doc.add_page_break()

    for title, rows in [
        ("Relapse Classification Map", CLASSIFICATION_MAP),
        ("Low-Capacity Version", LOW_CAPACITY),
        ("Ledger Entry Template", LEDGER_TEMPLATE),
        ("24-Hour Re-Entry Plan", REENTRY_PLAN),
        ("Pattern Review", PATTERN_REVIEW),
    ]:
        docx_paragraph(doc, title, size=18, color=DARK_PURPLE, bold=True)
        if title == "Relapse Classification Map":
            docx_paragraph(doc, "Use this map to identify which part of the self-mastery system failed to hold. More than one category may apply.")
        if title == "Low-Capacity Version":
            docx_paragraph(doc, "Use this version when the full ledger would increase pressure instead of restoring command.")
        if title == "Ledger Entry Template":
            docx_paragraph(doc, "Complete one entry for each meaningful breach. Keep language factual, brief, and specific.")
        if title == "24-Hour Re-Entry Plan":
            docx_paragraph(doc, "The next 24 hours should reduce repeat collapse. This page keeps the repair practical.")
        if title == "Pattern Review":
            docx_paragraph(doc, "Review ledger entries weekly during the Self-Mastery Blueprint and again at completion.")
        add_docx_table(doc, rows)
        if title == "24-Hour Re-Entry Plan":
            docx_paragraph(doc, "One-Sentence Re-Entry Declaration", size=13, color=DARK_PURPLE, bold=True, before=8)
            docx_paragraph(doc, "I am returning to command by protecting __________, completing __________, and updating __________.")
        if title != "Pattern Review":
            doc.add_page_break()

    docx_paragraph(doc, "Integration With the Self-Mastery Blueprint", size=13, color=DARK_PURPLE, bold=True, before=10)
    docx_paragraph(doc, "The Self-Mastery Blueprint is complete when the client can maintain, review, and restore her operating system without depending on constant external pressure. This ledger supports that standard by making disruption usable.")
    docx_paragraph(doc, "A mature self-mastery system does not require perfection. It requires accurate signals, honest classification, protected standards, and a reliable re-entry path.")
    doc.add_page_break()
    docx_paragraph(doc, "Blank Ledger Page", size=18, color=DARK_PURPLE, bold=True)
    add_docx_table(doc, [["Field", "Entry"], ["Date and Context", ""], ["Observed Breach", ""], ["System Area Affected", ""], ["Trigger or Friction", ""], ["Re-Entry Action", ""], ["Architecture Update", ""], ["Evidence After Re-Entry", ""]])

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle("kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=13, textColor=colors.HexColor(GOLD), alignment=TA_CENTER, spaceAfter=14),
        "title": ParagraphStyle("title", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=27, leading=31, textColor=colors.HexColor(DARK_PURPLE), alignment=TA_CENTER, spaceAfter=8),
        "subtitle": ParagraphStyle("subtitle", parent=styles["Normal"], fontName="Helvetica", fontSize=12, leading=16, textColor=colors.HexColor(SLATE), alignment=TA_CENTER, spaceAfter=20),
        "h1": ParagraphStyle("h1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor(DARK_PURPLE), spaceBefore=8, spaceAfter=8, keepWithNext=True),
        "h2": ParagraphStyle("h2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=16, textColor=colors.HexColor(DARK_PURPLE), spaceBefore=8, spaceAfter=5, keepWithNext=True),
        "body": ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica", fontSize=10.25, leading=13.2, textColor=colors.HexColor(BLACK), spaceAfter=6),
        "small": ParagraphStyle("small", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor(SLATE), spaceAfter=4),
        "cell": ParagraphStyle("cell", parent=styles["Normal"], fontName="Helvetica", fontSize=7.7, leading=9.4, textColor=colors.HexColor(BLACK), alignment=TA_LEFT),
        "cell_header": ParagraphStyle("cell_header", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7.5, leading=9.2, textColor=colors.white, alignment=TA_LEFT),
        "callout": ParagraphStyle("callout", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=13, textColor=colors.HexColor(DARK_PURPLE), spaceAfter=4),
    }


def paragraphize_table(rows, styles):
    converted = []
    for r, row in enumerate(rows):
        style = styles["cell_header"] if r == 0 else styles["cell"]
        converted.append([Paragraph(str(value), style) for value in row])
    return converted


def pdf_table(rows, styles, col_widths):
    table = LongTable(paragraphize_table(rows, styles), colWidths=col_widths, repeatRows=1, hAlign="LEFT", splitByRow=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(DARK_PURPLE)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(LIGHT_RULE)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
            ]
        )
    )
    return table


def on_page(canvas, doc):
    canvas.saveState()
    width, height = letter
    page = canvas.getPageNumber()
    if page > 1:
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor(DARK_PURPLE))
        canvas.drawRightString(width - 0.62 * inch, height - 0.35 * inch, "Distinct Character | Self-Mastery Blueprint")
        canvas.setFillColor(colors.HexColor(SLATE))
        canvas.drawCentredString(width / 2, 0.38 * inch, f"Copyright 2026 A. Solenne Institute. All rights reserved. | Page {page}")
    canvas.restoreState()


def add_pdf_bullets(story, items, styles):
    for item in items:
        story.append(Paragraph(f"- {item}", styles["body"]))


def add_pdf_numbered(story, items, styles):
    for index, item in enumerate(items, start=1):
        story.append(Paragraph(f"{index}. {item}", styles["body"]))


def add_section(story, title, body, styles):
    story.append(Paragraph(title, styles["h1"]))
    for text in body:
        story.append(Paragraph(text, styles["body"]))


def create_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.62 * inch,
        leftMargin=0.62 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    story = []

    if LOGO_PATH.exists():
        logo = Image(str(LOGO_PATH), width=0.9 * inch, height=0.9 * inch)
        logo.hAlign = "CENTER"
        story.append(Spacer(1, 0.35 * inch))
        story.append(logo)
        story.append(Spacer(1, 0.18 * inch))
    else:
        story.append(Spacer(1, 1 * inch))
    story.append(Paragraph("DISTINCT CHARACTER", styles["kicker"]))
    story.append(Paragraph(CONTENT["title"], styles["title"]))
    story.append(Paragraph(CONTENT["subtitle"], styles["subtitle"]))
    cover_rule = Table([[""]], colWidths=[2.2 * inch], rowHeights=[1])
    cover_rule.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(GOLD)), ("LINEABOVE", (0, 0), (-1, -1), 0, colors.HexColor(GOLD))]))
    cover_rule.hAlign = "CENTER"
    story.append(cover_rule)
    story.append(Spacer(1, 0.28 * inch))
    story.append(Paragraph("Designed for post-disruption review, capstone integration, and the preservation of self-mastery under real human conditions.", styles["subtitle"]))
    story.append(Spacer(1, 0.28 * inch))
    story.append(Paragraph("Self-Mastery Blueprint Companion Resource", styles["kicker"]))
    story.append(Paragraph("Version 1.0 | 2026", styles["small"]))
    story.append(PageBreak())

    add_section(
        story,
        "Copyright and Use Notice",
        [
            "The Relapse & Re-Entry Ledger is a proprietary educational resource within the Distinct Character Self-Mastery Blueprint. It is intended for personal protocol use, practitioner-supported review, and approved Distinct Character implementation contexts.",
            "No portion of this document may be copied, distributed, modified, taught, licensed, sold, or repackaged without written permission from A. Solenne Institute.",
            "Copyright 2026 A. Solenne Institute. All rights reserved. Distinct Character and related protocol language are part of the Distinct Character ecosystem.",
        ],
        styles,
    )
    story.append(Paragraph("Educational and Medical Disclaimer", styles["h2"]))
    for text in [
        "This document is educational and reflective in nature. It does not provide medical advice, mental health diagnosis, crisis care, or individualized treatment. It is not a substitute for care from a qualified medical, mental health, legal, or financial professional.",
        "If you are experiencing danger, self-harm thoughts, severe distress, medical symptoms, abuse, coercion, or inability to maintain safety, stop using this resource and seek immediate support from emergency services, crisis resources, or a qualified professional.",
        "Use low-capacity adaptations when the full process is too much. Re-entry should restore stability and function, not intensify pressure.",
    ]:
        story.append(Paragraph(text, styles["body"]))
    story.append(PageBreak())

    add_section(story, "Purpose", CONTENT["purpose"], styles)
    story.append(Paragraph("Core Distinction", styles["h2"]))
    story.append(pdf_table(CORE_DISTINCTIONS, styles, [0.9 * inch, 3.0 * inch, 2.5 * inch]))
    story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph("Use Rules", styles["h2"]))
    add_pdf_bullets(story, CONTENT["use_rules"], styles)
    story.append(PageBreak())

    story.append(Paragraph("Relapse Classification Map", styles["h1"]))
    story.append(Paragraph("Use this map to identify which part of the self-mastery system failed to hold. More than one category may apply.", styles["body"]))
    story.append(pdf_table(CLASSIFICATION_MAP, styles, [1.12 * inch, 1.92 * inch, 1.96 * inch, 1.4 * inch]))
    story.append(PageBreak())

    story.append(Paragraph("Immediate Re-Entry Sequence", styles["h1"]))
    story.append(Paragraph("Use this sequence before analysis becomes too large. The first re-entry goal is not total repair. It is restoration of command.", styles["body"]))
    add_pdf_numbered(story, CONTENT["reentry_steps"], styles)
    story.append(Paragraph("Low-Capacity Version", styles["h2"]))
    story.append(Paragraph("Use this version when the full ledger would increase pressure instead of restoring command.", styles["body"]))
    story.append(pdf_table(LOW_CAPACITY, styles, [2.4 * inch, 4.0 * inch]))
    story.append(PageBreak())

    story.append(Paragraph("Ledger Entry Template", styles["h1"]))
    story.append(Paragraph("Complete one entry for each meaningful breach. Keep language factual, brief, and specific.", styles["body"]))
    story.append(pdf_table(LEDGER_TEMPLATE, styles, [1.65 * inch, 4.75 * inch]))
    story.append(PageBreak())

    story.append(Paragraph("24-Hour Re-Entry Plan", styles["h1"]))
    story.append(Paragraph("The next 24 hours should reduce repeat collapse. This page keeps the repair practical.", styles["body"]))
    story.append(pdf_table(REENTRY_PLAN, styles, [1.0 * inch, 1.8 * inch, 1.8 * inch, 1.8 * inch]))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("One-Sentence Re-Entry Declaration", styles["h2"]))
    story.append(Paragraph("I am returning to command by protecting __________, completing __________, and updating __________.", styles["body"]))
    story.append(PageBreak())

    story.append(Paragraph("Pattern Review", styles["h1"]))
    story.append(Paragraph("Review ledger entries weekly during the Self-Mastery Blueprint and again at completion. The goal is to identify repeated system conditions, not to build a record of failure.", styles["body"]))
    story.append(pdf_table(PATTERN_REVIEW, styles, [2.25 * inch, 2.1 * inch, 2.05 * inch]))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("Integration With the Self-Mastery Blueprint", styles["h2"]))
    story.append(Paragraph("The Self-Mastery Blueprint is complete when the client can maintain, review, and restore her operating system without depending on constant external pressure. This ledger supports that standard by making disruption usable.", styles["body"]))
    story.append(Paragraph("A mature self-mastery system does not require perfection. It requires accurate signals, honest classification, protected standards, and a reliable re-entry path.", styles["body"]))
    story.append(PageBreak())

    story.append(Paragraph("Blank Ledger Page", styles["h1"]))
    story.append(Paragraph("Use this page for repeated entries or practitioner-guided review notes.", styles["body"]))
    story.append(pdf_table([["Field", "Entry"], ["Date and Context", ""], ["Observed Breach", ""], ["System Area Affected", ""], ["Trigger or Friction", ""], ["Re-Entry Action", ""], ["Architecture Update", ""], ["Evidence After Re-Entry", ""]], styles, [1.7 * inch, 4.7 * inch]))

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_docx()
    create_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
